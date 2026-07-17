from fastapi import FastAPI, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, ConfigDict
from typing import Optional, List
from matplotlib import colors
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import io
import markdown
import uuid
import json
from bs4 import BeautifulSoup

app = FastAPI()

api_key = os.environ.get("GROQ_API_KEY")
if api_key:
    from groq import Groq
    client = Groq(api_key=api_key)
else:
    client = None
    print("⚠️ GROQ_API_KEY not set. Word generator service will use fallback generation.")

# ===============================
# Request Model
# ===============================

class WordRequest(BaseModel):
    topic: Optional[str] = "Document"
    content: Optional[str] = None
    font_name: Optional[str] = "Times New Roman"
    font_size: Optional[int] = 12
    font_color: Optional[List[int]] = [0, 0, 0]
    font_color_name: Optional[str] = None
    heading_size: Optional[int] = 16
    pages: Optional[int] = 1
    user_input: Optional[str] = Field("", alias="message")

    model_config = ConfigDict(
        validate_by_name=True,
        extra="ignore"
    )
def color_to_rgb(color_input):
    """Convert any color name or hex to RGBColor"""
    try:
        if isinstance(color_input, list) and len(color_input) == 3:
            return RGBColor(color_input[0], color_input[1], color_input[2])
        elif isinstance(color_input, str):
            r, g, b = colors.to_rgb(color_input)
            return RGBColor(int(r*255), int(g*255), int(b*255))
        else:
            return RGBColor(0, 0, 0)
    except:
        return RGBColor(0, 0, 0)

# ===============================
# AI Content Generator - FIXED
# ===============================

WORDS_PER_PAGE = 300

def generate_ai_content(topic: str, pages: int = 1):
    try:
        pages = max(1, min(pages, 50))
        words_needed = pages * WORDS_PER_PAGE
        
        # Much more specific prompt to force actual content generation
        prompt = f"""You are an academic writer. Write a detailed, well-structured document about "{topic}".

This is a direct instruction - you MUST write the actual content now, not ask for more information.

The document should be approximately {words_needed} words (around {pages} pages).

Write in markdown format with:
# {topic} (as the main title)

## Abstract
Write a 100-150 word summary.

## Introduction
Introduce the topic, its importance, and what the document will cover.

## Main Sections
Create at least 3 major sections with ## headings, and include subsections with ### headings where appropriate.
Fill each section with substantial, factual content about {topic}.

## Conclusion
Summarize key points and suggest future directions.

## References
Include 5-7 academic-style references.

Start writing the document now. Do not ask for confirmation. Do not explain what you're going to do. Just write the document content directly.
"""
        
        print(f"📝 Generating {pages}-page document on '{topic}'...")
        
        if client is None:
            print("⚠️ No Groq client available; using local fallback document generator.")
            return f"# {topic}\n\n## Abstract\nThis document provides an overview of {topic}.\n\n## Introduction\n{topic} is an important subject that deserves comprehensive analysis.\n\n## Main Sections\n\n### Section 1\nThis section covers the main concepts and context of {topic}.\n\n### Section 2\nHere we explain key details and explain why {topic} matters.\n\n### Section 3\nThis section explores outcomes, ideas, and next steps.\n\n## Conclusion\n{topic} continues to evolve with important implications for future work.\n\n## References\n1. Example Reference 1\n2. Example Reference 2\n"  

        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a professional academic writer. You always write complete documents without asking for confirmation."},
                {"role": "user", "content": prompt}
            ],
            model="openai/gpt-oss-20b",
            temperature=0.7,
            max_tokens=4000
        )

        generated_text = response.choices[0].message.content
        word_count = len(generated_text.split())
        print(f"✅ Generated {word_count} words (target was {words_needed})")
        
        # If the response is asking for more info, generate a fallback
        if "need more information" in generated_text.lower() or "could you please" in generated_text.lower():
            print("⚠️ Got question instead of content, using fallback generation")
            fallback_prompt = f"""Write a comprehensive document about "{topic}".
            
Start with # {topic}
Then write an abstract, introduction, at least 3 detailed sections, a conclusion, and references.
Make it approximately {words_needed} words.
Just write the content directly without any questions."""
            
            fallback_response = client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are a professional writer. Write the document now without questions."},
                    {"role": "user", "content": fallback_prompt}
                ],
                model="openai/gpt-oss-20b",
                temperature=0.7,
                max_tokens=4000
            )
            generated_text = fallback_response.choices[0].message.content
        
        return generated_text
    
    except Exception as e:
        print(f"❌ AI generation error: {str(e)}")
        # Return a basic document instead of failing
        return f"""## Abstract
This document provides an overview of {topic}.

## Introduction
{topic} is an important subject that deserves comprehensive analysis.

## Section 1: Overview
This section covers the fundamental aspects of {topic}.

## Section 2: Key Developments
Recent developments in {topic} have been significant.

## Section 3: Future Directions
The future of {topic} holds many possibilities.

## Conclusion
In conclusion, {topic} continues to evolve and impact various fields.

## References
1. Author, A. (2023). Understanding {topic}. Academic Press.
2. Author, B. (2024). Recent Advances in {topic}. Journal of Studies.
"""

def normalize_markdown_content(markdown_text, topic):
    """Remove a duplicate leading H1 title if it matches the final topic."""
    if not markdown_text:
        return markdown_text

    lines = markdown_text.splitlines()
    if lines:
        first_line = lines[0].strip()
        normalized_topic = topic.strip().strip('"')
        heading_variants = [f"# {normalized_topic}", f"# \"{normalized_topic}\"", f"# The Topic Of \"{normalized_topic}\""
                            ]
        if first_line in heading_variants:
            # Remove the first heading and any immediately following blank line
            lines = lines[1:]
            if lines and not lines[0].strip():
                lines = lines[1:]

    return "\n".join(lines)


def convert_markdown_to_docx(document, markdown_text, font_name, font_size, 
                            title_color, section_color, all_headings_color, body_color):
    """Convert markdown to formatted DOCX with different colors for different heading levels"""
    
    html = markdown.markdown(markdown_text, extensions=['tables'])
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.find_all(recursive=False):
        # Headings
        if element.name in ["h1", "h2", "h3", "h4"]:
            level = int(element.name[1])
            heading = document.add_heading(element.text, level=level)
            
            # Determine which color to use for this heading
            if level == 1:  # Main title (H1)
                if title_color:
                    heading_rgb = title_color
                elif all_headings_color:
                    heading_rgb = all_headings_color
                else:
                    heading_rgb = body_color
                print(f"   H{level} (title) using color: {heading_rgb}")
            else:  # Subheadings (H2, H3, H4)
                if section_color:
                    heading_rgb = section_color
                elif all_headings_color:
                    heading_rgb = all_headings_color
                else:
                    heading_rgb = body_color
                print(f"   H{level} (section) using color: {heading_rgb}")
            
            for run in heading.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size + 4 - level)
                run.font.color.rgb = heading_rgb
                if level == 1:
                    run.bold = True

        # Paragraphs (body text - always body_color)
        elif element.name == "p":
            paragraph = document.add_paragraph()
            for content in element.contents:
                if hasattr(content, 'get_text'):
                    text = content.get_text()
                else:
                    text = str(content)
                
                if text.strip():
                    run = paragraph.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = body_color

                    if hasattr(content, 'name'):
                        if content.name in ['strong', 'b']:
                            run.bold = True
                        if content.name in ['em', 'i']:
                            run.italic = True

        # Bullet List
        elif element.name == "ul":
            for li in element.find_all("li"):
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(li.text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = body_color

        # Numbered List
        elif element.name == "ol":
            for li in element.find_all("li"):
                p = document.add_paragraph(style="List Number")
                run = p.add_run(li.text)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = body_color

        # Tables
        elif element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue

            header_cells = rows[0].find_all(["td", "th"])
            table = document.add_table(rows=len(rows), cols=len(header_cells))
            table.style = "Table Grid"

            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    cell_text = cell.text.strip()
                    table_cell = table.cell(i, j)
                    
                    table_cell.paragraphs[0].clear()
                    run = table_cell.paragraphs[0].add_run(cell_text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = body_color

                    if i == 0:
                        run.bold = True

# ===============================
# Parser Function
# ===============================

def normalize_topic_text(topic: str) -> str:
    """Normalize a parsed topic by stripping quotes and topic prefixes."""
    if not topic:
        return topic

    topic = topic.strip()
    topic = re.sub(r'^(?:the\s+)?topic(?:\s+of)?\s*[:\-–—]?\s*', '', topic, flags=re.IGNORECASE)
    topic = topic.strip(' "\'')
    topic = re.sub(r'\s+', ' ', topic)
    return topic


def parse_user_request(user_input):
    """Parse the user's request into document generation parameters."""

    if client is None:
        print("⚠️ No Groq client available; using local parse fallback.")
        return local_parse_user_request(user_input)

    parse_prompt = f"""You are an AI assistant that extracts document information from user requests.

User request: "{user_input}"

Extract the following and return ONLY a valid JSON object:
- topic: The main subject ONLY - IMPORTANT: Return this with each word capitalized (e.g., "Artificial Intelligence", not "artificial intelligence")
- pages: Number of pages (integer, default 1)
- font_name: Font name (string, default "Times New Roman")
- font_size: Font size (integer, default 12)
- title_color: Color for the main title ONLY (string, null if not specified)
- section_color: Color for section headings ONLY (string, null if not specified)
- all_headings_color: Color for all headings if one color is specified for all (string, null if not specified)

Examples:
Input: "write a document on growth of ai with title color blue, section heading color light blue"
Output: {{"topic": "Growth Of Ai", "pages": 1, "font_name": "Times New Roman", "font_size": 12, "title_color": "blue", "section_color": "light blue", "all_headings_color": null}}

Input: "create a 10 page document about climate change with arial font, heading color red"
Output: {{"topic": "Climate Change", "pages": 10, "font_name": "arial", "font_size": 12, "title_color": null, "section_color": null, "all_headings_color": "red"}}

Return ONLY the JSON object, no other text.
"""

    try:
        parse_response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a precise JSON extractor. Return only valid JSON with capitalized topic words."},
                {"role": "user", "content": parse_prompt}
            ],
            model="openai/gpt-oss-20b",
            temperature=0.1,
            max_tokens=500
        )
        
        result_text = parse_response.choices[0].message.content
        print(f"🧠 Groq parse response: {result_text}")
        
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            # Normalize and capitalize topic
            if parsed.get('topic'):
                normalized_topic = normalize_topic_text(parsed['topic'])
                words = normalized_topic.split()
                parsed['topic'] = ' '.join([word if word.isupper() else word.capitalize() for word in words])
            return parsed
    except Exception as e:
        print(f"⚠️ Parsing failed: {e}")
        return local_parse_user_request(user_input)
    
    return local_parse_user_request(user_input)


def local_parse_user_request(user_input):
    """Fallback parser for document requests when Groq is unavailable."""
    parsed = {
        "topic": "Document",
        "pages": 1,
        "font_name": "Times New Roman",
        "font_size": 12,
        "title_color": None,
        "section_color": None,
        "all_headings_color": None
    }

    if not user_input:
        return parsed

    # Extract pages
    pages_match = re.search(r'(\d+)\s*pages?', user_input, re.IGNORECASE)
    if pages_match:
        parsed["pages"] = int(pages_match.group(1))

    # Extract font size
    font_size_match = re.search(r'(?:font\s*size|size)\s*[:=]?\s*(\d+)', user_input, re.IGNORECASE)
    if font_size_match:
        parsed["font_size"] = int(font_size_match.group(1))

    # Extract font name
    font_name_match = re.search(r'font\s*[:=]?\s*([a-zA-Z ]+)', user_input, re.IGNORECASE)
    if font_name_match:
        font_name = font_name_match.group(1).strip()
        if font_name.lower() not in ["size", "color", "pages", "page"]:
            parsed["font_name"] = font_name.title()

    # Extract colors
    title_color_match = re.search(r'title\s*color\s*[:=]?\s*([a-zA-Z ]+)', user_input, re.IGNORECASE)
    if title_color_match:
        parsed["title_color"] = title_color_match.group(1).strip().lower()

    section_color_match = re.search(r'section\s*heading\s*color\s*[:=]?\s*([a-zA-Z ]+)', user_input, re.IGNORECASE)
    if section_color_match:
        parsed["section_color"] = section_color_match.group(1).strip().lower()

    all_headings_color_match = re.search(r'heading\s*color\s*[:=]?\s*([a-zA-Z ]+)', user_input, re.IGNORECASE)
    if all_headings_color_match and not parsed["section_color"]:
        parsed["all_headings_color"] = all_headings_color_match.group(1).strip().lower()

    # Extract topic after keywords like topic, about, on, or topic of
    topic_match = re.search(
        r'(?:topic|about|on)(?:\s+of)?\s+["\']?(.+?)(?=(?:\s+(?:with|using|for|by|from|and|that|which|where|while|font|size|pages|page|document|file|color|heading|title|section)\b|$))',
        user_input,
        re.IGNORECASE
    )
    if topic_match:
        topic = topic_match.group(1).strip()
    else:
        # Remove common instruction prefixes and trailing formatting instructions
        topic = re.sub(
            r'^(?:write|create|make|generate|produce)\s+(?:a|an)?\s*(?:word\s+)?(?:document|file|report)?\s*(?:on|about|for)?\s*(.*)',
            r'\1',
            user_input,
            flags=re.IGNORECASE
        ).strip()
        topic = re.sub(
            r'\s+(?:with|using|for|by|from|and|that|which|where|while|font|size|pages|page|document|file|color|heading|title|section)\b.*$',
            '',
            topic,
            flags=re.IGNORECASE
        ).strip()

    if topic:
        topic = normalize_topic_text(topic)
        words = topic.split()
        parsed["topic"] = ' '.join([
            word if word.isupper() else word.capitalize()
            for word in words
        ])

    return parsed

# ===============================
# Main Endpoint
# ===============================

@app.post("/generate-word")
def generate_word(request: WordRequest, background_tasks: BackgroundTasks):
    try:
        document = Document()
        
        # Default values
        title_color = None
        section_color = None
        all_headings_color = None
        final_topic = request.topic
        
        # Parse user input with Groq
        if request.user_input:
            parsed = parse_user_request(request.user_input)
            
            if parsed:
                if parsed.get('topic'):
                    final_topic = parsed['topic']
                    print(f"📚 Topic: {final_topic}")
                
                if parsed.get('pages'):
                    request.pages = int(parsed['pages'])
                    print(f"📄 Pages: {request.pages}")
                
                if parsed.get('font_name'):
                    request.font_name = parsed['font_name'].title()
                    print(f"✒️ Font: {request.font_name}")
                
                if parsed.get('font_size'):
                    request.font_size = int(parsed['font_size'])
                    print(f"🔤 Font size: {request.font_size}")
                
                title_color = parsed.get('title_color')
                section_color = parsed.get('section_color')
                all_headings_color = parsed.get('all_headings_color')
                
                print(f"🎨 Title color: {title_color}")
                print(f"🎨 Section color: {section_color}")
                print(f"🎨 All headings color: {all_headings_color}")

        # Generate content
        content = request.content or generate_ai_content(final_topic, request.pages)

        # Add main title
        heading = document.add_heading(final_topic, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set body color to black
        body_color = RGBColor(0, 0, 0)
        
        # Convert color names to RGB
        title_rgb = color_to_rgb(title_color) if title_color else None
        section_rgb = color_to_rgb(section_color) if section_color else None
        all_headings_rgb = color_to_rgb(all_headings_color) if all_headings_color else None
        
        print(f"\n🎨 Final colors:")
        print(f"   Title: {title_rgb}")
        print(f"   Sections: {section_rgb}")
        print(f"   All headings: {all_headings_rgb}")
        
        # Apply color to main title
        for run in heading.runs:
            run.font.name = request.font_name
            run.font.size = Pt(request.heading_size)
            run.bold = True
            if title_rgb:
                run.font.color.rgb = title_rgb
            elif all_headings_rgb:
                run.font.color.rgb = all_headings_rgb

        # Normalize markdown so duplicate H1 title is not converted twice
        content = normalize_markdown_content(content, final_topic)

        # Convert markdown to docx
        convert_markdown_to_docx(
            document,
            content,
            request.font_name,
            request.font_size,
            title_rgb,
            section_rgb,
            all_headings_rgb,
            body_color
        )

        # Save file
        os.makedirs("generated_files", exist_ok=True)
        file_name = f"{uuid.uuid4()}.docx"
        file_path = os.path.join("generated_files", file_name)
        document.save(file_path)
        
        background_tasks.add_task(os.remove, file_path)

        # Create filename from topic
        safe_topic = re.sub(r'[^\w\s-]', '', final_topic)
        safe_topic = re.sub(r'[-\s]+', '_', safe_topic)
        safe_topic = re.sub(r'_+', '_', safe_topic).strip('_')
        download_name = f"{final_topic}.docx"  # Keep original with spaces

        print(f"✅ Document saved as: {download_name}")
        
        return FileResponse(
            path=file_path,
            filename=download_name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")
    
@app.post("/parse-word-request")
def parse_word_request(request: WordRequest):
    """Just parse the request without generating the document"""
    try:
        parsed_params = {
            "topic": request.topic,
            "pages": request.pages,
            "font_name": request.font_name,
            "font_size": request.font_size,
            "title_color": None,
            "section_color": None,
            "all_headings_color": None
        }
        
        if request.user_input:
            parsed = parse_user_request(request.user_input)
            
            if parsed:
                if parsed.get('topic'):
                    parsed_params["topic"] = parsed['topic']
                if parsed.get('pages'):
                    parsed_params["pages"] = int(parsed['pages'])
                if parsed.get('font_name'):
                    parsed_params["font_name"] = parsed['font_name'].title()
                if parsed.get('font_size'):
                    parsed_params["font_size"] = int(parsed['font_size'])
                if parsed.get('title_color'):
                    parsed_params["title_color"] = parsed['title_color']
                if parsed.get('section_color'):
                    parsed_params["section_color"] = parsed['section_color']
                if parsed.get('all_headings_color'):
                    parsed_params["all_headings_color"] = parsed['all_headings_color']
        
        return parsed_params
        
    except Exception as e:
        print(f"❌ Error parsing request: {str(e)}")
        return {
            "topic": request.topic or "Document",
            "pages": request.pages,
            "font_name": request.font_name,
            "font_size": request.font_size,
            "title_color": None,
            "section_color": None,
            "all_headings_color": None
        }


def create_word_document_bytes(user_input: str):
    """Create a .docx file in memory from a user request string."""
    try:
        request = WordRequest(message=user_input)
    except Exception:
        request = WordRequest()

    final_topic = request.topic or "Document"
    parsed = None
    if request.user_input:
        parsed = parse_user_request(request.user_input)

    if parsed:
        if parsed.get('topic'):
            final_topic = parsed['topic']
        if parsed.get('pages'):
            request.pages = int(parsed['pages'])
        if parsed.get('font_name'):
            request.font_name = parsed['font_name'].title()
        if parsed.get('font_size'):
            request.font_size = int(parsed['font_size'])

    content = generate_ai_content(final_topic, request.pages)
    content = normalize_markdown_content(content, final_topic)

    document = Document()
    heading = document.add_heading(final_topic, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in heading.runs:
        run.font.name = request.font_name
        run.font.size = Pt(request.heading_size)
        run.bold = True

    body_color = RGBColor(0, 0, 0)
    title_rgb = None
    section_rgb = None
    all_headings_rgb = None

    convert_markdown_to_docx(
        document,
        content,
        request.font_name,
        request.font_size,
        title_rgb,
        section_rgb,
        all_headings_rgb,
        body_color
    )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    safe_topic = re.sub(r'[^\w\s-]', '', final_topic)
    safe_topic = re.sub(r'[-\s]+', '_', safe_topic)
    safe_topic = re.sub(r'_+', '_', safe_topic).strip('_')
    download_name = f"{safe_topic or 'document'}.docx"

    return buffer.read(), download_name


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "word_generator_agent:app",
        host="0.0.0.0",
        port=8000,
        reload=True
    )